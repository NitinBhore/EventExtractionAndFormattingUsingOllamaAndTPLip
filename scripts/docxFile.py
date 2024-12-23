from langchain_community.llms import Ollama
from docx import Document
import json
import re

# Initialize the Ollama model
llm = Ollama(model='llama3.2:latest')

# Function to read content from a .docx file
def read_docx(file_path):
    doc = Document(file_path)
    full_text = []

    for para in doc.paragraphs:
        full_text.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return "\n".join(full_text)

# Function to extract events using the LLM
def extract_events_with_llm(content):
    prompt = f"""
    Extract events from the following content:
    {content}

    Return the extracted events in the following JSON format:
    {{
        "events": [
            [
                "<start_date>",
                "<day_of_week>",
                [
                    [
                        "<repeat_count>",
                        "<start_time>",
                        "<end_time>",
                        "<event_type>",
                        "<setup_style>",
                        "<number_of_people>",
                        "<additional_comments>"
                    ]
                ]
            ]
        ]
    }}
    """

    try:
        response = llm.invoke(prompt)
        # print("LLM Response:", response)  # Debugging: Print the response

        # Use regex to find and extract the JSON part
        json_match = re.search(r'\{.*\}', response, re.DOTALL)
        if json_match:
            json_content = json_match.group(0)  # Extract the matched JSON part
            
            # Fix the JSON to escape quotes correctly
            json_content = json_content.replace('“', '"').replace('”', '"').replace('\'', '"')

            return json.loads(json_content)  # Attempt to parse the extracted JSON

        print("No valid JSON found in the response.")
        return {"events": []}

    except json.JSONDecodeError as e:
        print("Failed to decode JSON from LLM response:", e)
        # print("Response content:", response)  # Print the response for debugging
        return {"events": []}
    except Exception as e:
        print("An error occurred while extracting events:", e)
        return {"events": []}

# Function to write the output to a JSON file
def write_json_to_file(data, filename):
    with open(filename, 'w') as json_file:
        json.dump(data, json_file, indent=2)

# Main function to read, process, and return output
def docs_data_extraction(docx_file_path):
    # Step 1: Read the .docx file
    content = read_docx(docx_file_path)
    # print("Document Content:\n", content)  # Debugging: Print the entire content

    # Step 2: Extract events using the LLM
    events = extract_events_with_llm(content)

    # Step 3: Write the extracted events to a JSON file
    output_filename = "output/extracted_events.json"
    write_json_to_file(events, output_filename)
    print(f"\n\nData written to {output_filename}")

# Usage
# docx_file = 'RFP3.docx'  # Replace with your .docx file path
# docs_data_extraction(docx_file)

